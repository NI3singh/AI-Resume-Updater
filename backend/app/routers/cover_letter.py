"""Cover Letter AI.

Generates a job/company/résumé-tailored cover letter that reuses the user's exact
Word template (`app/templates/cover_letter_template.docx`): the layout, fonts,
colours, header and contact block are preserved by editing the .docx in place —
only the body cell (salutation → opening → JD bridge → 3 résumé-grounded bullets
→ closing → sign-off) is rebuilt from AI content. Download as DOCX always, or PDF
via headless LibreOffice when available.

Two steps:
  POST /cover-letter/generate  -> structured, résumé-grounded letter content (LLM)
  POST /cover-letter/render    -> fill the template -> DOCX (or PDF) bytes
"""

import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, Response, status

from ..config import settings
from ..deps import get_current_user
from ..llm import LLMError, LLMNotConfigured, chat_json
from ..models import User
from ..schemas import (
    CoverLetterContent,
    CoverLetterGenerateIn,
    CoverLetterGenerateOut,
    CoverLetterRenderIn,
)

router = APIRouter(prefix="/cover-letter", tags=["cover-letter"])

_TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "cover_letter_template.docx"
# Formatting captured from the template.
_BODY_PT = 12      # body paragraphs + bullets
_HEAD_PT = 14      # salutation, closing block (sign-off / name / signature)
_SIG_FONT = "Courgette"  # the cursive signature line

_DIGIT_RE = re.compile(r"\d[\d,.]*")
# Inline formatting: **bold** and __underline__ (the app's conventions).
_FMT_RE = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__)")


def _digits(text: str) -> set[str]:
    """Normalized numeric tokens — to flag numbers not grounded in the résumé/JD."""
    return {m.replace(",", "").rstrip(".") for m in _DIGIT_RE.findall(text or "")}


COVER_LETTER_SYSTEM = (
    "You are an elite cover-letter writer. Using ONLY the candidate's résumé plus "
    "the target job description (JD), company, and position, write a tailored, "
    "professional cover letter. Mirror this exact structure:\n"
    "- salutation: 'Dear {Company} Hiring Team,' (use the real company; a polite "
    "generic greeting if unknown).\n"
    "- opening: ONE paragraph expressing genuine interest in the **{position}** "
    "position at {company}, grounded in the candidate's background (degree, area, "
    "years of experience as shown in the résumé).\n"
    "- bridge: ONE sentence of the form 'Your job description emphasizes ... :' "
    "naming 1-2 key needs from the JD.\n"
    "- bullets: EXACTLY 3 bullets. Each begins with a **bold lead-in:** then a "
    "sentence that ties a SPECIFIC résumé project, skill, or experience to the JD's "
    "needs. Name the real project/technology from the résumé (e.g. a project title).\n"
    "- closing: 2 short paragraphs — one company-specific (why this company), one "
    "forward-looking.\n"
    "- signoff: 'Yours Sincerely'.\n\n"
    "STRICT RULES:\n"
    "- Use ONLY facts, projects, skills, employers, and numbers present in the "
    "résumé. NEVER invent metrics, employers, tools, dates, or outcomes. Use JD "
    "keywords ONLY where the résumé genuinely supports them.\n"
    "- ~250-350 words, confident and specific — not generic filler.\n"
    "- Mark key skills/terms and the bullet lead-ins with **double asterisks** for "
    "bold. Do NOT leave placeholders like [Company]; fill them in.\n"
    "- Output MUST be ONE valid JSON object — no markdown, no code fences, no text "
    "before or after. Inside string values NEVER use the double-quote character "
    "(write names plainly or in 'single quotes'); never put a line break inside a "
    "string value.\n\n"
    "JSON shape:\n"
    "{\n"
    '  "headerTitle": "a short professional title for the candidate (from the résumé)",\n'
    '  "salutation": "Dear ... Hiring Team,",\n'
    '  "opening": "...",\n'
    '  "bridge": "Your job description emphasizes ... :",\n'
    '  "bullets": ["**Lead-in:** ...", "**Lead-in:** ...", "**Lead-in:** ..."],\n'
    '  "closing": ["...", "..."],\n'
    '  "signoff": "Yours Sincerely",\n'
    '  "signatureName": "the candidate\'s name",\n'
    '  "missing_keywords": ["JD requirements the résumé does not support"]\n'
    "}"
)


# ── Résumé → grounding text ───────────────────────────────────────────────────
def _resume_text(data: dict[str, Any]) -> str:
    """Compact, readable serialization of the résumé for the LLM (and number guard)."""
    parts: list[str] = []
    personal = data.get("personal") if isinstance(data.get("personal"), dict) else {}
    if personal.get("name"):
        parts.append(f"Name: {personal['name']}")
    if personal.get("summary"):
        parts.append(f"Summary: {personal['summary']}")

    def lst(key: str) -> list[dict]:
        v = data.get(key)
        return [e for e in v if isinstance(e, dict)] if isinstance(v, list) else []

    skills = []
    for c in lst("skills"):
        items = ", ".join(str(i) for i in (c.get("items") or []) if isinstance(i, str))
        if c.get("category") or items:
            skills.append(f"{c.get('category', '')}: {items}".strip())
    if skills:
        parts.append("Skills:\n" + "\n".join(skills))

    def block(key: str, head_fields: list[str]) -> list[str]:
        out = []
        for e in lst(key):
            head = " / ".join(str(e.get(f)) for f in head_fields if e.get(f))
            if key == "projects" and e.get("techStack"):
                head += f" ({e['techStack']})"
            bullets = "\n".join(
                "  - " + str(b) for b in (e.get("bullets") or []) if isinstance(b, str) and b.strip()
            )
            entry = (head + ("\n" + bullets if bullets else "")).strip()
            if entry:
                out.append(entry)
        return out

    exp = block("experience", ["role", "company"])
    if exp:
        parts.append("Experience:\n" + "\n".join(exp))
    proj = block("projects", ["name"])
    if proj:
        parts.append("Projects:\n" + "\n".join(proj))

    edu = []
    for e in lst("education"):
        edu.append(" / ".join(str(e.get(f)) for f in ("degree", "institution", "highlight") if e.get(f)))
    edu = [e for e in edu if e]
    if edu:
        parts.append("Education:\n" + "\n".join(edu))

    for key, label in (("achievements", "Achievements"), ("certifications", "Certifications")):
        items = [str(e.get("text")) for e in lst(key) if e.get("text")]
        if items:
            parts.append(label + ":\n" + "\n".join("  - " + i for i in items))

    return "\n\n".join(parts)[: settings.parse_text_limit]


# ── DOCX template fill ────────────────────────────────────────────────────────
def _run(para, text: str, size: int, bold: bool, underline: bool, font_name: str | None) -> None:
    r = para.add_run(text)
    r.font.size = Pt(size)
    if bold:
        r.bold = True
    if underline:
        r.underline = True
    if font_name:
        r.font.name = font_name


def _add_para(cell, text: str, size: int, *, align=None, bold_all: bool = False,
              underline_all: bool = False, font_name: str | None = None):
    """Add a paragraph, parsing **bold** and __underline__ spans into formatted
    runs; empty text => a blank line."""
    para = cell.add_paragraph()
    if align is not None:
        para.alignment = align
    if not text:
        _run(para, "", size, bold_all, underline_all, font_name)
        return para
    for seg in _FMT_RE.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            _run(para, seg[2:-2], size, True, underline_all, font_name)
        elif seg.startswith("__") and seg.endswith("__"):
            _run(para, seg[2:-2], size, bold_all, True, font_name)
        else:
            _run(para, seg, size, bold_all, underline_all, font_name)
    return para


def _set_paragraph_text(para, text: str) -> None:
    """Replace a paragraph's text while preserving its first run's formatting."""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def _find_body_cell(doc):
    """The body cell = the table cell holding the most text (the letter). Returns
    ``(cell, row)`` so the caller can normalize the row's forced height."""
    best, best_row, best_len = None, None, -1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                n = len(cell.text)
                if n > best_len:
                    best, best_row, best_len = cell, row, n
    return best, best_row


def _find_header(doc):
    """Locate the header name + tagline paragraphs (the cell with the biggest font)."""
    best_cell, best_sz = None, 0.0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue
                    for r in p.runs:
                        sz = r.font.size.pt if r.font.size else 0.0
                        if sz > best_sz:
                            best_sz, best_cell = sz, cell
    if best_cell is None:
        return None, None
    nonempty = [p for p in best_cell.paragraphs if p.text.strip()]
    name_p = nonempty[0] if nonempty else None
    tagline_p = nonempty[1] if len(nonempty) > 1 else None
    return name_p, tagline_p


def _fill_template(content: CoverLetterContent, name: str) -> bytes:
    if not _TEMPLATE.exists():
        raise HTTPException(status_code=500, detail="Cover letter template is missing on the server.")
    doc = Document(str(_TEMPLATE))

    # Header: refresh name + tagline (keep everything else — contact block, colours).
    name_p, tagline_p = _find_header(doc)
    if name_p is not None and name:
        _set_paragraph_text(name_p, name)
    if tagline_p is not None and content.headerTitle.strip():
        _set_paragraph_text(tagline_p, content.headerTitle.strip())

    # Body: rebuild from the generated content.
    cell, body_row = _find_body_cell(doc)
    if cell is None:
        raise HTTPException(status_code=500, detail="Cover letter template body cell not found.")
    # The template's body row ships with a huge forced height (~11in) that pushes
    # the letter onto a second page, plus a stray character in a spacer cell. Make
    # the row auto-size, clear spacer junk, and top-align the body so the letter
    # sits cleanly under the header rule.
    if body_row is not None:
        try:
            body_row.height = None
            body_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        except Exception:  # noqa: BLE001 — best-effort layout tidy-up
            pass
        for spacer in body_row.cells:
            # Compare the underlying <w:tc> — row.cells returns fresh wrappers and
            # the gridSpan'd body cell repeats, so object identity isn't reliable.
            if spacer._tc is not cell._tc and spacer.text.strip():
                spacer.text = ""
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    except Exception:  # noqa: BLE001
        pass
    _clear_cell(cell)
    left, right = WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT

    salutation = _add_para(cell, content.salutation.strip() or "Dear Hiring Team,", _HEAD_PT, align=left, bold_all=True)
    salutation.paragraph_format.space_before = Pt(24)  # consistent gap under the header rule
    _add_para(cell, "", _BODY_PT)
    if content.opening.strip():
        _add_para(cell, content.opening.strip(), _BODY_PT)
        _add_para(cell, "", _BODY_PT)
    if content.bridge.strip():
        _add_para(cell, content.bridge.strip(), _BODY_PT)
    for b in content.bullets:
        if isinstance(b, str) and b.strip():
            _add_para(cell, "- " + b.strip(), _BODY_PT)
    if any(isinstance(b, str) and b.strip() for b in content.bullets):
        _add_para(cell, "", _BODY_PT)
    for c in content.closing:
        if isinstance(c, str) and c.strip():
            _add_para(cell, c.strip(), _BODY_PT)
            _add_para(cell, "", _BODY_PT)
    _add_para(cell, content.signoff.strip() or "Yours Sincerely", _HEAD_PT, align=right)
    if name:
        _add_para(cell, name, _HEAD_PT, align=right, bold_all=True)
    sig = content.signatureName.strip() or name
    if sig:
        # The signature is the cursive Courgette name, underlined (matches the template).
        _add_para(cell, sig, _HEAD_PT, align=right, bold_all=True, underline_all=True, font_name=_SIG_FONT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── DOCX → PDF via headless LibreOffice ───────────────────────────────────────
def _find_soffice() -> str | None:
    if settings.soffice_path and Path(settings.soffice_path).exists():
        return settings.soffice_path
    for cand in ("soffice", "libreoffice", "soffice.exe"):
        found = shutil.which(cand)
        if found:
            return found
    for p in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/opt/libreoffice/program/soffice",
    ):
        if Path(p).exists():
            return p
    return None


def _docx_to_pdf(docx_bytes: bytes) -> bytes:
    soffice = _find_soffice()
    if not soffice:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="PDF needs LibreOffice on the server. Your DOCX downloaded fine — open it and "
                   "Save as PDF, or install LibreOffice (or set SOFFICE_PATH) to enable PDF here.",
        )
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "cover_letter.docx"
        src.write_bytes(docx_bytes)
        profile = (Path(tmp) / "loprofile").as_uri()
        try:
            subprocess.run(
                [soffice, f"-env:UserInstallation={profile}", "--headless",
                 "--convert-to", "pdf", "--outdir", tmp, str(src)],
                check=True, capture_output=True, timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as exc:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail="PDF conversion failed. The DOCX is still available to download.",
            ) from exc
        pdf = Path(tmp) / "cover_letter.pdf"
        if not pdf.exists():
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail="PDF conversion produced no file. The DOCX is still available to download.",
            )
        return pdf.read_bytes()


# ── Routes ───────────────────────────────────────────────────────────────────
@router.post("/generate", response_model=CoverLetterGenerateOut)
def generate(
    payload: CoverLetterGenerateIn, current_user: User = Depends(get_current_user),
) -> CoverLetterGenerateOut:
    """Generate the structured, résumé-grounded letter content for the given JD."""
    jd = payload.job_description[: settings.parse_text_limit]
    data = payload.data if isinstance(payload.data, dict) else {}
    resume = _resume_text(data)
    if not resume.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Your résumé is empty — add some content before generating a cover letter.",
        )
    company = payload.company.strip()
    job_title = payload.job_title.strip()
    instruction = payload.instruction.strip()[: settings.parse_text_limit]

    header = (
        f"COMPANY: {company or '(not provided)'}\n"
        f"POSITION / JOB TITLE: {job_title or '(infer from the JD)'}\n"
    )
    if instruction:
        header += f"\n*** USER CHANGE REQUEST (apply this; never invent facts): {instruction}\n"
    user_msg = header + "\nJOB DESCRIPTION:\n" + jd + "\n\n=== CANDIDATE RÉSUMÉ ===\n" + resume
    if instruction:
        user_msg += f"\n\nREMINDER — apply the user request above: {instruction}"

    try:
        result = chat_json(COVER_LETTER_SYSTEM, user_msg, max_tokens=2048, temperature=0.5)
    except LLMNotConfigured as exc:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=str(exc)) from exc

    personal = data.get("personal") if isinstance(data.get("personal"), dict) else {}
    resume_name = str(personal.get("name") or "").strip()

    def _s(key: str) -> str:
        v = result.get(key)
        return v.strip() if isinstance(v, str) else ""

    def _list(key: str) -> list[str]:
        v = result.get(key)
        return [x.strip() for x in v if isinstance(x, str) and x.strip()] if isinstance(v, list) else []

    content = CoverLetterContent(
        headerTitle=_s("headerTitle"),
        salutation=_s("salutation") or (f"Dear {company} Hiring Team," if company else "Dear Hiring Team,"),
        opening=_s("opening"),
        bridge=_s("bridge"),
        bullets=_list("bullets"),
        closing=_list("closing"),
        signoff=_s("signoff") or "Yours Sincerely",
        signatureName=_s("signatureName") or resume_name,
    )

    # Anti-fabrication: flag (don't strip) any number not grounded in résumé/JD.
    allowed = _digits(resume) | _digits(jd) | _digits(company) | _digits(job_title)
    warnings: list[str] = []
    checks = [("Opening", content.opening), ("Bridge", content.bridge)]
    checks += [(f"Bullet {i + 1}", b) for i, b in enumerate(content.bullets)]
    checks += [(f"Closing paragraph {i + 1}", c) for i, c in enumerate(content.closing)]
    for label, text in checks:
        extra = _digits(text) - allowed
        if extra:
            warnings.append(
                f"{label}: has number(s) not found in your résumé/JD "
                f"({', '.join(sorted(extra))}) — double-check before sending."
            )

    return CoverLetterGenerateOut(content=content, warnings=warnings, missing_keywords=_list("missing_keywords")[:20])


@router.post("/render")
def render(payload: CoverLetterRenderIn, current_user: User = Depends(get_current_user)) -> Response:
    """Fill the template with the (edited) content and return DOCX or PDF bytes."""
    docx_bytes = _fill_template(payload.content, payload.name.strip())
    safe = re.sub(r"[^\w-]+", "_", payload.name.strip()).strip("_")[:80] or "Cover_Letter"
    if payload.format.lower() == "pdf":
        pdf = _docx_to_pdf(docx_bytes)
        return Response(
            content=pdf, media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe}_Cover_Letter.pdf"'},
        )
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe}_Cover_Letter.docx"'},
    )
