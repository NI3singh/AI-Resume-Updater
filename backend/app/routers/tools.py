"""Tool routes — resume import pipeline, JD tailoring + LaTeX compilation.

All endpoints require an authenticated user.
- POST /tools/extract-text : uploaded PDF/DOCX/TXT -> plain text (+ hyperlinks)
- POST /tools/parse        : extracted text        -> structured ResumeData (LLM)
- POST /tools/verify       : text + parsed data     -> corrected data + warnings (LLM)
- POST /tools/transform/plan    : resume + JD        -> ordered list of units to tailor
- POST /tools/transform/section : one unit + JD (+ optional sources) -> guarded rewrite
- POST /tools/compile      : LaTeX source           -> compiled PDF bytes

Hyperlinks: in a PDF the URL behind a link is stored as a *link annotation*
(and in DOCX as a relationship), not in the text layer — plain text extraction
only yields the anchor text ("Portfolio", "[View Credential]", "Github Link",
...). The extractors therefore also collect the real URLs and append them to
the returned text as a "HYPERLINKS" block, and the parse/verify prompts teach
the LLM how to map each URL back onto the right field.
"""

import io
import json

import httpx
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, HTTPException, Response, UploadFile, status
from pypdf import PdfReader

from ..config import settings
from ..deps import get_current_user
from ..llm import LLMError, LLMNotConfigured, chat_json
from ..models import User
from ..schemas import (
    CompileIn,
    ExtractOut,
    ParseIn,
    ParseOut,
    TransformPlanIn,
    TransformPlanOut,
    TransformSectionIn,
    TransformSectionOut,
    TransformStep,
    VerifyIn,
    VerifyOut,
    VerifySummary,
)
from ..transform_guard import guard_unit, prepare_for_llm

router = APIRouter(prefix="/tools", tags=["tools"])

# Top-level sections the parser targets (must match the frontend ResumeData type).
_SECTION_KEYS = [
    "personal", "education", "skills", "projects", "experience",
    "extracurricular", "achievements", "certifications", "publications",
]

# Shape the LLM must produce. Custom (user-defined) sections are intentionally
# excluded — too freeform to infer — and are preserved on the client.
_SCHEMA = """{
  "personal": { "name": "", "email": "", "phone": "", "location": "", "linkedin": "", "github": "", "website": "", "summary": "" },
  "education": [ { "institution": "", "degree": "", "location": "", "gpaFormat": "CGPA | GPA | Percentage | Grade | (empty)", "gpaLabel": "", "startDate": "", "endDate": "", "highlight": "" } ],
  "skills": [ { "category": "", "items": ["", ""] } ],
  "projects": [ { "name": "", "techStack": "", "githubUrl": "", "liveUrl": "", "bullets": ["", ""] } ],
  "experience": [ { "role": "", "company": "", "location": "", "startDate": "", "endDate": "", "current": false, "projectSubtitle": "", "bullets": ["", ""] } ],
  "extracurricular": [ { "title": "", "organization": "", "startDate": "", "endDate": "", "bullets": ["", ""] } ],
  "achievements": [ { "text": "" } ],
  "certifications": [ { "text": "", "credentialUrl": "" } ],
  "publications": [ { "title": "", "authors": "", "abstractText": "", "paperUrl": "", "paperLinkLabel": "" } ]
}"""

_RULES = """Rules:
- Return ONLY a single JSON object with exactly these top-level keys: personal, education, skills, projects, experience, extracurricular, achievements, certifications, publications. No markdown, no code fences, no commentary.
- NEVER invent information. If something is not in the source, use "" for text and [] for lists. Do not fabricate emails, phone numbers, URLs, GPAs, or dates.
- Dates must be formatted as "MMM YYYY" (e.g., "Jan 2023"). For an ongoing role/study use "Present" as endDate; for an ongoing job also set "current": true.
- "bullets" and "items" are arrays of strings (one entry each). Keep wording faithful to the source — do not embellish or summarize away detail.
- education.gpaFormat: choose from what the source shows — "8.9/10" -> CGPA (gpaLabel "8.9"), "3.8/4" -> GPA (gpaLabel "3.8"), "85%" -> Percentage (gpaLabel "85"), "A+" -> Grade (gpaLabel "A+"); if unclear, leave gpaFormat "" and put the raw value in gpaLabel.
- Do NOT include any "id" fields. Leave a section's array empty if you cannot populate it.
- personal.summary: include a professional summary/objective only if the resume actually has one."""

# Shared by PARSE_SYSTEM and VERIFY_SYSTEM: how to use the appended URL list.
_LINKS_RULE = """Hyperlink rule:
- The resume text may end with a "=== HYPERLINKS EXTRACTED FROM THE DOCUMENT ===" block. In PDFs/DOCX the URL behind a link is invisible to text extraction — the body above shows only anchor text such as "Portfolio", "[View Credential]", "Github Link", "Live Link", or "arXiv Pre-print, 2025". The block lists the real URLs in reading order (top of the document to the bottom); entries may also look like "anchor text -> URL".
- Use these URLs (and any URLs written literally in the body) to fill the URL fields. Never invent or guess a URL.
  * mailto: address -> personal.email (drop the "mailto:" prefix)
  * linkedin.com URL -> personal.linkedin
  * github.com/<user> profile URL (no repository path) -> personal.github
  * a personal/portfolio website -> personal.website
  * github.com/<user>/<repo> -> the githubUrl of the project whose name best matches the repo name
  * a deployed-app URL (e.g. *.vercel.app, *.netlify.app, *.onrender.com, or a custom domain) appearing among the project links -> that project's liveUrl
  * credential URLs (credly, coursera, udemy, ...) -> certifications[].credentialUrl, matched to the certifications in order
  * arxiv.org / doi.org / journal URLs -> publications[].paperUrl, with paperLinkLabel set to the visible anchor text in the body (e.g. "arXiv Pre-print, 2025")
- The block's order matches the order the links appear in the resume — use it to disambiguate repeated anchors such as several "[View Credential]" bullets.
- Always output complete URLs including the scheme (https://...)."""

PARSE_SYSTEM = (
    "You are a precise resume parser. Convert the resume text the user provides "
    "into structured JSON matching exactly this schema:\n\n"
    + _SCHEMA
    + "\n\n"
    + _RULES
    + "\n\n"
    + _LINKS_RULE
)

VERIFY_SYSTEM = (
    "You are a meticulous resume-data QA reviewer. You are given the ORIGINAL "
    "resume text and a PARSED JSON produced from it. Your job:\n"
    "1. Correct values that don't match the source (typos, wrong dates, mis-split "
    "bullets, fields placed in the wrong section).\n"
    "2. Fill in values clearly present in the text but missing from the JSON.\n"
    "3. Normalize dates to \"MMM YYYY\"/\"Present\" and fix gpaFormat/gpaLabel.\n"
    "4. NEVER invent data that isn't in the source.\n"
    "5. Apply the hyperlink rule below: every URL field (linkedin, github, website, "
    "githubUrl, liveUrl, credentialUrl, paperUrl) must be filled from the HYPERLINKS "
    "block when one is present, and corrected if mismatched.\n\n"
    + _LINKS_RULE
    + "\n\nThe corrected resume must follow exactly this schema:\n\n"
    + _SCHEMA
    + "\n\nReturn ONLY a JSON object of this form:\n"
    "{\n"
    '  "data": <the corrected resume object>,\n'
    '  "warnings": ["short notes about anything missing, ambiguous, or low-confidence"],\n'
    '  "summary": { "sections_found": [section keys that have content], "missing": [important fields/sections that appear absent] }\n'
    "}\nNo markdown, no code fences, no commentary."
)

# NOTE: TRANSFORM_SYSTEM is the legacy one-shot whole-resume prompt. The live
# app no longer uses it (replaced by the interactive per-unit flow below), but
# the evaluation harness (backend/eval/run_differential.py) imports it to study
# the raw-vs-guarded behaviour — keep it defined.
TRANSFORM_SYSTEM = (
    "You are a resume tailoring assistant. You are given a candidate's resume as "
    "structured JSON — every entry carries an \"id\" field — and a target job "
    "description (JD). Tailor the resume to the job using ONLY what is already in it.\n\n"
    "ALLOWED (this is the whole job):\n"
    "- Rephrase experience/project/extracurricular \"bullets\", personal.summary, and "
    "education \"highlight\" to foreground what matters for this JD.\n"
    "- Reorder entries within a section and items within skills so the most relevant "
    "come first.\n"
    "- Drop entries, bullets, or skill items that are clearly irrelevant to this job "
    "(omit them from your output).\n"
    "- Use the JD's terminology in rephrased text ONLY where the resume genuinely "
    "demonstrates that experience (e.g. say \"REST APIs\" instead of \"web endpoints\" "
    "if they built them — but never claim a tool/skill the resume doesn't show).\n\n"
    "FORBIDDEN — copy these through verbatim, character for character:\n"
    "- Every \"id\" — each entry you output must keep the exact id it had. NEVER "
    "create a new entry.\n"
    "- personal: name, email, phone, location, linkedin, github, website.\n"
    "- experience: role, company, location, startDate, endDate, current, projectSubtitle.\n"
    "- education: institution, degree, location, gpaFormat, gpaLabel, startDate, endDate.\n"
    "- projects: name, techStack, githubUrl, liveUrl. skills: category names.\n"
    "- achievements, certifications, publications: every field (you may only reorder "
    "or drop whole entries).\n"
    "- Every URL anywhere. Every number, percentage, metric, and date — NEVER write a "
    "number that does not appear in that same entry in the source resume.\n\n"
    "The resume JSON follows this schema (plus the \"id\" field on every entry):\n\n"
    + _SCHEMA
    + "\n\nReturn ONLY a JSON object of this form:\n"
    "{\n"
    '  "data": <the tailored resume object, same schema, ids preserved>,\n'
    '  "changes": ["short human-readable notes on what you changed and why"],\n'
    '  "match": {\n'
    '    "covered_keywords": [JD skills/requirements the resume genuinely demonstrates],\n'
    '    "missing_keywords": [JD requirements the resume does NOT show - be honest, do NOT add them to the resume]\n'
    "  }\n"
    "}\nNo markdown, no code fences, no commentary."
)

# Interactive Transform: tailor ONE unit at a time. The model only ever emits
# the small editable piece (bullets / a text field / the summary) — never the
# whole resume and never immutable facts — which removes both the truncation and
# the verbatim-echo corruption of the old one-shot transform.
SECTION_SYSTEM = (
    "You are a precise resume bullet editor. You are given ONE unit of a "
    "candidate's resume and a target job description (JD). Rephrase ONLY the "
    "editable text of this unit to foreground what matters for the JD.\n\n"
    "STRICT RULES:\n"
    "- Use ONLY facts already present in the provided unit and in any SOURCES the "
    "user pasted (e.g. a project README). NEVER invent skills, tools, employers, "
    "outcomes, metrics, or dates.\n"
    "- NEVER introduce a number, percentage, metric, or date that is not in the "
    "unit or the SOURCES.\n"
    "- You may reorder, merge, sharpen, or drop a bullet that is irrelevant to "
    "this JD. Use the JD's terminology only where the unit genuinely shows it.\n"
    "- Do NOT restate names, companies, roles, dates, URLs, or tech-stack names — "
    "those are fixed elsewhere; edit only the wording of the content.\n"
    "- Keep each bullet truthful, specific, and concise.\n"
    "- A USER CHANGE REQUEST (when present) is your TOP priority for HOW to "
    "rewrite — length, wording, emphasis, terminology, ordering. You MUST visibly "
    "apply it and produce a genuinely DIFFERENT result from before. The only "
    "limit is truth: never invent facts or numbers to satisfy it; if it asks for "
    "something the unit and SOURCES don't support, apply the rest and omit that part.\n\n"
    "Return ONLY a JSON object (no markdown, no commentary) with the ONE content "
    "key that matches the unit, plus the metadata keys:\n"
    "{\n"
    '  "bullets": ["..."],   // experience / projects / extracurricular units\n'
    '  "text": "...",        // an education "highlight" unit\n'
    '  "summary": "...",     // the personal summary unit\n'
    '  "items": ["..."],     // a skills unit (reorder/trim only - no new items)\n'
    '  "rationale": "one short sentence on what you changed and why",\n'
    '  "covered_keywords": ["JD terms this unit genuinely demonstrates"],\n'
    '  "missing_keywords": ["JD terms this unit does NOT show - do not add them"],\n'
    '  "no_change_recommended": false\n'
    "}\n"
    "Include only the single content key for this unit's kind."
)


def _pdf_link_urls(reader: PdfReader) -> list[str]:
    """Collect hyperlink URLs from PDF link annotations, in reading order.

    URLs live in each page's /Annots (entries with /Subtype /Link and an /A
    action carrying /URI) and are invisible to ``page.extract_text()``. Links
    are sorted top-to-bottom, then left-to-right (the PDF y-axis points up).
    Best-effort: malformed annotations are skipped silently.
    """
    urls: list[str] = []
    seen: set[str] = set()
    for page in reader.pages:
        try:
            annots = page.get("/Annots")
            if not annots:
                continue
            positioned: list[tuple[float, float, str]] = []
            for ref in annots.get_object():
                try:
                    annot = ref.get_object()
                    if annot.get("/Subtype") != "/Link":
                        continue
                    action = annot.get("/A")
                    uri = action.get_object().get("/URI") if action is not None else None
                    if not uri:
                        continue
                    x0 = y0 = 0.0
                    rect = annot.get("/Rect")
                    if rect is not None:
                        rect = rect.get_object()
                        if len(rect) == 4:
                            x0, y0 = float(rect[0]), float(rect[1])
                    positioned.append((-y0, x0, str(uri).strip()))
                except Exception:
                    continue
            for _, _, uri in sorted(positioned):
                if uri and uri not in seen:
                    seen.add(uri)
                    urls.append(uri)
        except Exception:
            continue
    return urls


def _extract_pdf(raw: bytes) -> tuple[str, list[str]]:
    reader = PdfReader(io.BytesIO(raw))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return text, _pdf_link_urls(reader)


def _extract_docx(raw: bytes) -> tuple[str, list[str]]:
    doc = DocxDocument(io.BytesIO(raw))
    parts: list[str] = []
    links: list[str] = []
    seen: set[str] = set()

    def collect_links(paragraph) -> None:
        # python-docx >= 1.1 exposes hyperlinks together with their anchor text.
        for h in getattr(paragraph, "hyperlinks", []):
            addr = (h.address or "").strip()
            if not addr or addr in seen:
                continue
            seen.add(addr)
            anchor = (h.text or "").strip()
            links.append(f"{anchor} -> {addr}" if anchor else addr)

    for p in doc.paragraphs:
        parts.append(p.text)
        collect_links(p)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
            for cell in row.cells:
                for p in cell.paragraphs:
                    collect_links(p)
    return "\n".join(parts), links


def _links_block(links: list[str]) -> str:
    """Render the collected URLs as a block the LLM can map back onto fields."""
    if not links:
        return ""
    listed = "\n".join(f"- {u}" for u in links[:80])  # cap to keep the prompt small
    return (
        "\n\n=== HYPERLINKS EXTRACTED FROM THE DOCUMENT (reading order) ===\n"
        "(The resume text above shows only the anchor text of each link — "
        "these are the real URLs behind them.)\n" + listed
    )


def _sections_present(data: dict) -> list[str]:
    """Builtin section keys that actually carry content (for the review summary)."""
    found: list[str] = []
    for key in _SECTION_KEYS:
        value = data.get(key)
        if key == "personal":
            if isinstance(value, dict) and any(str(v).strip() for v in value.values()):
                found.append(key)
        elif isinstance(value, list) and value:
            found.append(key)
    return found


@router.post("/extract-text", response_model=ExtractOut)
def extract_text(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
) -> ExtractOut:
    raw = file.file.read()
    if len(raw) > settings.upload_max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large (max {settings.upload_max_bytes // (1024 * 1024)} MB).",
        )
    if not raw:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="The file is empty.")

    name = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()
    links: list[str] = []
    try:
        if name.endswith(".pdf") or "pdf" in ctype:
            text, links = _extract_pdf(raw)
        elif name.endswith(".docx") or "wordprocessingml" in ctype:
            text, links = _extract_docx(raw)
        elif name.endswith(".txt") or ctype.startswith("text/"):
            text = raw.decode("utf-8", errors="ignore")
        elif name.endswith(".doc"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Legacy .doc isn't supported — save it as PDF or DOCX and try again.",
            )
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported file type. Please upload a PDF, DOCX, or TXT file.",
            )
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001 — corrupt/unreadable document
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not read the document: {exc}",
        ) from exc

    text = (text or "").strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No selectable text found. Scanned or image-only PDFs aren't supported.",
        )
    # Truncate the body FIRST, then append the hyperlink block, so the URLs are
    # never cut off by the length limit on long resumes.
    return ExtractOut(text=text[: settings.parse_text_limit] + _links_block(links))


@router.post("/parse", response_model=ParseOut)
def parse_resume(payload: ParseIn, current_user: User = Depends(get_current_user)) -> ParseOut:
    try:
        data = chat_json(PARSE_SYSTEM, payload.text)
    except LLMNotConfigured as exc:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=str(exc)) from exc
    return ParseOut(data=data)


@router.post("/verify", response_model=VerifyOut)
def verify_resume(payload: VerifyIn, current_user: User = Depends(get_current_user)) -> VerifyOut:
    user_msg = (
        "ORIGINAL RESUME TEXT:\n"
        + payload.text
        + "\n\nPARSED JSON TO REVIEW:\n"
        + json.dumps(payload.data, ensure_ascii=False)
    )
    try:
        result = chat_json(VERIFY_SYSTEM, user_msg)
    except LLMNotConfigured as exc:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail=str(exc)) from exc
    except LLMError:
        # Resilient: never block the user — keep the parsed data and flag it.
        return VerifyOut(
            data=payload.data,
            warnings=["Verification step was unavailable — please review the imported fields carefully."],
            summary=VerifySummary(sections_found=_sections_present(payload.data)),
        )

    data = result.get("data")
    if not isinstance(data, dict):
        data = payload.data

    warnings_raw = result.get("warnings") or []
    warnings = (
        [str(w) for w in warnings_raw]
        if isinstance(warnings_raw, list)
        else [str(warnings_raw)]
    )

    summary_raw = result.get("summary") if isinstance(result.get("summary"), dict) else {}
    sections_found = summary_raw.get("sections_found")
    missing = summary_raw.get("missing")
    summary = VerifySummary(
        sections_found=(
            [str(s) for s in sections_found]
            if isinstance(sections_found, list)
            else _sections_present(data)
        ),
        missing=[str(m) for m in missing] if isinstance(missing, list) else [],
    )
    return VerifyOut(data=data, warnings=warnings, summary=summary)


# Units the interactive Transform can tailor (summary + the text-bearing sections).
_TAILORABLE_KINDS = {"summary", "experience", "projects", "extracurricular", "education", "skills"}


def _clean_keywords(value: object, limit: int = 20) -> list[str]:
    """De-duplicated, trimmed, capped list of keyword strings from LLM output."""
    if not isinstance(value, list):
        return []
    out: list[str] = []
    for k in value:
        s = str(k).strip()
        if s and s not in out:
            out.append(s)
        if len(out) >= limit:
            break
    return out


# Plan enrichment: a small metadata-only LLM pass that flags which units are
# worth tailoring for this JD (so the wizard can suggest "skip" for ones that
# already fit) and reports the résumé's honest gaps. No rewritten text here, so
# it's cheap and can't truncate.
PLAN_SYSTEM = (
    "You are a resume tailoring planner. You are given a target job description "
    "(JD) and a list of resume UNITS — each with an index, kind, label, and a "
    "short preview of its current content. For each unit, decide whether it is "
    "worth tailoring for THIS job.\n\n"
    "Return ONLY a JSON object (no markdown, no commentary):\n"
    "{\n"
    '  "units": [ { "i": <index>, "recommend_change": true, "reason": "<= 12 words why" } ],\n'
    '  "missing_keywords": ["JD requirements the resume as a whole does not show"]\n'
    "}\n"
    "Set recommend_change=false when a unit is already well aligned with the JD, "
    "or is irrelevant to it, so the user can safely skip it. Keep each reason "
    "short and concrete. Be honest in missing_keywords and NEVER suggest adding "
    "them to the resume."
)


def _unit_preview(kind: str, entry: dict, limit: int = 200) -> str:
    """A short, single-line snapshot of a unit's current content for the planner."""
    if kind == "summary":
        text = str(entry.get("summary") or "")
    elif kind == "education":
        text = str(entry.get("highlight") or "")
    else:
        bullets = [str(b) for b in (entry.get("bullets") or []) if isinstance(b, (str, int, float))]
        text = " · ".join(bullets)
    return " ".join(text.split())[:limit]


def _enrich_plan(
    steps: list[TransformStep], previews: list[str], jd: str, job_title: str, company: str,
) -> list[str]:
    """Best-effort: ask the LLM which units to tailor + the JD gaps. Mutates
    ``steps`` in place (recommend_change/reason) and returns missing_keywords.
    Any LLM failure leaves the deterministic plan (all recommend_change=True)."""
    if not steps:
        return []
    units = [
        {"i": k, "kind": s.kind, "label": s.label, "preview": previews[k]}
        for k, s in enumerate(steps)
    ]
    user_msg = (
        f"JOB TITLE: {job_title.strip() or '(not provided)'}\n"
        f"COMPANY: {company.strip() or '(not provided)'}\n\n"
        "JOB DESCRIPTION:\n" + jd + "\n\n"
        "RESUME UNITS (JSON):\n" + json.dumps(units, ensure_ascii=False)
    )
    try:
        result = chat_json(PLAN_SYSTEM, user_msg, max_tokens=2048)
    except (LLMError, LLMNotConfigured):
        return []  # planning is optional — the deterministic plan still works

    raw_units = result.get("units")
    if isinstance(raw_units, list):
        for u in raw_units:
            if not isinstance(u, dict):
                continue
            try:
                idx = int(u.get("i"))
            except (TypeError, ValueError):
                continue
            if 0 <= idx < len(steps):
                steps[idx].recommend_change = bool(u.get("recommend_change", True))
                reason = str(u.get("reason") or "").strip()
                if reason:
                    steps[idx].reason = reason[:140]
    return _clean_keywords(result.get("missing_keywords"))


@router.post("/transform/plan", response_model=TransformPlanOut)
def transform_plan(
    payload: TransformPlanIn, current_user: User = Depends(get_current_user),
) -> TransformPlanOut:
    """Return the ordered list of résumé units the wizard will tailor, one at a
    time, annotated with JD relevance. The actual rewriting happens per-unit in
    /transform/section."""
    present = _sections_present(payload.data)
    if not present or present == ["personal"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Your resume is empty — add some content before tailoring it.",
        )

    prepared = prepare_for_llm(payload.data)
    steps: list[TransformStep] = []
    previews: list[str] = []

    personal = prepared.get("personal") or {}
    summary = personal.get("summary")
    if isinstance(summary, str) and summary.strip():
        steps.append(TransformStep(
            kind="summary", entry_id="", section="personal", label="Professional Summary",
        ))
        previews.append(_unit_preview("summary", personal))

    for section, label_field, noun in [
        ("experience", "company", "Experience"),
        ("projects", "name", "Project"),
        ("extracurricular", "title", "Activity"),
    ]:
        for e in prepared.get(section) or []:
            if not isinstance(e, dict):
                continue
            label = str(e.get(label_field) or "").strip() or noun
            steps.append(TransformStep(
                kind=section, entry_id=str(e.get("id") or ""), section=section, label=label,
                asks_readme=(section == "projects"),
                asks_related_work=(section == "experience"),
            ))
            previews.append(_unit_preview(section, e))

    for e in prepared.get("education") or []:
        if isinstance(e, dict) and str(e.get("highlight") or "").strip():
            label = str(e.get("institution") or "").strip() or "Education"
            steps.append(TransformStep(
                kind="education", entry_id=str(e.get("id") or ""), section="education", label=label,
            ))
            previews.append(_unit_preview("education", e))

    missing = _enrich_plan(
        steps, previews, payload.job_description[: settings.parse_text_limit],
        payload.job_title, payload.company,
    )
    return TransformPlanOut(steps=steps, missing_keywords=missing)


@router.post("/transform/section", response_model=TransformSectionOut)
def transform_section(
    payload: TransformSectionIn, current_user: User = Depends(get_current_user),
) -> TransformSectionOut:
    """Tailor ONE résumé unit to the JD. The model only sees (and only returns)
    this unit's editable text, so it can't truncate or corrupt the rest of the
    résumé; the guard then strips any number not present in the unit or in the
    user-provided sources (e.g. a pasted README)."""
    if payload.kind not in _TAILORABLE_KINDS:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unknown unit kind.")

    jd = payload.job_description[: settings.parse_text_limit]
    sources = [
        s[: settings.parse_text_limit]
        for s in payload.sources
        if isinstance(s, str) and s.strip()
    ]
    instruction = payload.instruction.strip()[: settings.parse_text_limit]

    # The change request is placed BOTH up top and at the very end — for the
    # summary step the grounding "SOURCES" block is the whole résumé, so an
    # instruction tucked only at the bottom would get buried.
    header = (
        f"JOB TITLE: {payload.job_title.strip() or '(not provided)'}\n"
        f"COMPANY: {payload.company.strip() or '(not provided)'}\n"
        f"UNIT KIND: {payload.kind}\n"
    )
    if instruction:
        header += (
            "\n*** USER CHANGE REQUEST (top priority — you MUST apply this and "
            f"return a different result, without inventing facts): {instruction}\n"
        )

    user_msg = (
        header
        + "\nJOB DESCRIPTION:\n" + jd
        + "\n\nTHIS RESUME UNIT (JSON):\n" + json.dumps(payload.entry, ensure_ascii=False)
    )
    if sources:
        user_msg += (
            "\n\nSOURCES the user provided — facts and numbers appearing here are "
            "allowed in the rewrite:\n" + "\n\n---\n\n".join(sources)
        )
    if instruction:
        user_msg += f"\n\nREMINDER — apply the user's change request above: {instruction}"

    # A little randomness so a regeneration (especially with a change request)
    # yields a genuinely different, instruction-following result rather than the
    # identical output a reasoning model returns at temperature 0.
    try:
        result = chat_json(SECTION_SYSTEM, user_msg, max_tokens=4096, temperature=0.6)
    except LLMNotConfigured as exc:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=str(exc)) from exc

    guarded, warnings = guard_unit(payload.kind, payload.entry, result, sources=sources)
    return TransformSectionOut(
        proposal=guarded,
        rationale=str(result.get("rationale") or "").strip(),
        covered_keywords=_clean_keywords(result.get("covered_keywords")),
        missing_keywords=_clean_keywords(result.get("missing_keywords")),
        warnings=warnings,
        no_change_recommended=bool(result.get("no_change_recommended")),
    )


@router.post("/compile")
def compile_pdf(payload: CompileIn, current_user: User = Depends(get_current_user)) -> Response:
    pdf: bytes | None = None
    try:
        with httpx.Client(timeout=60.0) as client:
            primary = client.post(
                settings.latex_compile_url,
                json={
                    "compiler": "pdflatex",
                    "resources": [{"main": True, "content": payload.latex}],
                },
            )
            if primary.status_code < 400 and primary.content:
                pdf = primary.content
            else:
                # Fallback service (GET with the LaTeX as a query param).
                alt = client.get(settings.latex_compile_fallback, params={"text": payload.latex})
                if alt.status_code < 400 and alt.content:
                    pdf = alt.content
    except httpx.HTTPError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Compilation service error: {exc}",
        ) from exc

    if not pdf:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="LaTeX compilation failed. Check your LaTeX code for errors.",
        )

    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="resume.pdf"'},
    )